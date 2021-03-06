from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import docx
import requests
import io
import re
import tempfile
import base64

from utils.common import get_valid_xml_string as xstr


def _write_file(r, fp):
    for chunk in r.iter_content(chunk_size=1024):
        if chunk:
            fp.write(chunk)
    return fp


def _first_child_found_in(parent, tagnames):
    """
    Return the first child of parent with tag in *tagnames*, or None if
    not found.
    """
    for tagname in tagnames:
        child = parent.find(qn(tagname))
        if child is not None:
            return child
    return None


def _insert_element_before(parent, elm, successors):
    """
    Insert *elm* as child of *parent* before any existing child having
    tag name found in *successors*.
    """
    successor = _first_child_found_in(parent, successors)
    if successor is not None:
        successor.addprevious(elm)
    else:
        parent.append(elm)
    return elm


class Run:
    """
    Single run inside a paragraph
    """
    def __init__(self, ref):
        self.ref = ref

    def add_text(self, text):
        self.ref.add_text(xstr(text))
        return self

    def add_image(self, image):
        if image and len(image) > 0:
            fimage = tempfile.NamedTemporaryFile()
            if re.search(r'http[s]?://', image):
                image = requests.get(image, stream=True)
                _write_file(image, fimage)
            else:
                image = base64.b64decode(image.split(',')[1])
                fimage.write(image)
            self.ref.add_picture(fimage)


class Paragraph:
    """
    One paragraph: supports normal text runs, hyperlinks,
    horizontal lines.
    """
    def __init__(self, ref):
        self.ref = ref

    def add_run(self, text=None):
        return Run(self.ref.add_run(xstr(text)))

    def add_hyperlink(self, url, text):
        part = self.ref.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        new_run = docx.oxml.shared.OxmlElement('w:r')
        r_pr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(r_pr)
        new_run.text = xstr(text)

        hyperlink.append(new_run)
        self.ref._p.append(hyperlink)

        r = docx.text.run.Run(new_run, self.ref)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return self

    def add_horizontal_line(self):
        p = self.ref._p
        p_pr = p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')

        _insert_element_before(p_pr, p_bdr, successors=(
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct',
            'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
            'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
            'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        ))

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        p_bdr.append(bottom)

        return self

    def justify(self):
        self.ref.paragraph_format.alignment = \
            docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        return self


class Document:
    """
    A docx document representation
    """
    def __init__(self, template=None):
        self.doc = docx.Document(template)

    def add_paragraph(self, text=None):
        return Paragraph(self.doc.add_paragraph(xstr(text)))

    def add_image(self, image):
        if image and len(image):
            sec = self.doc.sections[-1]
            try:
                cols = int(sec._sectPr.xpath('./w:cols')[0].get(qn('w:num')))
                width = ((sec.page_width / cols) -
                         (sec.right_margin + sec.left_margin))
            except Exception:
                width = (sec.page_width - (sec.right_margin + sec.left_margin))

            fimage = tempfile.NamedTemporaryFile()
            if re.search(r'http[s]?://', image):
                image = requests.get(image, stream=True)
                _write_file(image, fimage)
            else:
                image = base64.b64decode(image.split(',')[1])
                fimage.write(image)
            self.doc.add_picture(fimage, width=width)
        return self

    def add_heading(self, text, level):
        self.doc.add_heading(text, level=level)
        return self

    def add_page_break(self):
        self.doc.add_page_break()
        return self

    def save_to_file(self, fp):
        self.doc.save(fp)

    def save(self):
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()
